import streamlit as st
import pandas as pd
import numpy as np
import plotly as plt
import plotly.graph_objects as go
import tempfile
import os
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document
from docx.shared import Inches

# Page config
st.set_page_config(
    page_title="CRISAP - Climate Risk Assessment Platform",
    page_icon="🌍",
    layout="wide"
)

# CSS for better styling
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        color: #1E88E5;
    }
    .sub-header {
        font-size: 1.8rem;
        color: #0D47A1;
    }
    .info-box {
        background-color: #E3F2FD;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 5px solid #1E88E5;
    }
    .warning-box {
        background-color: #FFF8E1;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 5px solid #FFC107;
    }
    .success-text {
        color: #2E7D32;
        font-weight: bold;
    }
</style>
""", unsafe_allow_html=True)

# App title
st.markdown("<h1 class='main-header'>🚀 CRISAP - Global Climate Risk Dashboard</h1>", unsafe_allow_html=True)
st.markdown("*Climate Risk and Impact Scenarios Analysis Platform*")

# Sidebar for controls
with st.sidebar:
    st.image("https://via.placeholder.com/150x80", caption="CRISAP v1.0")
    st.markdown("### 🌍 Region Selection")
    
    # Country & Region Selection with more countries
    country_list = ["Ghana", "Nigeria", "Kenya", "South Africa", "Egypt", "Morocco", "Tanzania", "Ethiopia"]
    selected_country = st.selectbox("Select Country", country_list)
    
    region_list = {
        "Ghana": ["Greater Accra", "Ashanti", "Northern", "Western", "Eastern", "Central"],
        "Nigeria": ["Lagos", "Abuja", "Kano", "Rivers", "Kaduna", "Oyo"],
        "Kenya": ["Nairobi", "Mombasa", "Kisumu", "Nakuru", "Eldoret"],
        "South Africa": ["Gauteng", "Western Cape", "KwaZulu-Natal", "Eastern Cape"],
        "Egypt": ["Cairo", "Alexandria", "Giza", "Luxor"],
        "Morocco": ["Casablanca", "Rabat", "Marrakech", "Fes"],
        "Tanzania": ["Dar es Salaam", "Zanzibar", "Arusha", "Mwanza"],
        "Ethiopia": ["Addis Ababa", "Dire Dawa", "Bahir Dar", "Mekelle"]
    }
    
    if selected_country in region_list:
        selected_region = st.selectbox("Select Region", region_list[selected_country])
    else:
        selected_region = "All Regions"
    
    # Time frame selector
    st.markdown("### ⏱️ Time Horizon")
    selected_timeframe = st.select_slider(
        "Select projection timeframe",
        options=["2030", "2050", "2070", "2100"]
    )
    
    # Climate scenario selection
    st.markdown("### 🌡️ Climate Scenario")
    selected_scenario = st.radio(
        "Select IPCC scenario",
        ["RCP 2.6 (Low emissions)", "RCP 4.5 (Intermediate)", "RCP 8.5 (High emissions)"]
    )
    
    # Analysis type
    st.markdown("### 📊 Analysis Type")
    analysis_types = st.multiselect(
        "Select impact categories",
        ["Economic", "Agricultural", "Water Resources", "Health", "Infrastructure"],
        default=["Economic", "Agricultural"]
    )

# Create tabs for different sections
tab1, tab2, tab3 = st.tabs(["Risk Assessment", "Data Visualization", "Report Generation"])

# Generate mock data based on selections
@st.cache_data
def generate_climate_data(country, region, timeframe, scenario):
    # Create a seed based on inputs for consistent random data
    seed = hash(f"{country}_{region}_{timeframe}_{scenario}") % 10000
    np.random.seed(seed)
    
    # Base risk values that change with region and scenario
    scenario_multiplier = {
        "RCP 2.6 (Low emissions)": 0.7,
        "RCP 4.5 (Intermediate)": 1.0,
        "RCP 8.5 (High emissions)": 1.5
    }
    
    time_multiplier = {
        "2030": 0.8,
        "2050": 1.0,
        "2070": 1.3,
        "2100": 1.6
    }
    
    # Base economic loss calculated with inputs and randomness
    base_economic_loss = len(region) * 2.5 * scenario_multiplier[scenario] * time_multiplier[timeframe]
    economic_loss = round(base_economic_loss * (1 + np.random.normal(0, 0.2)), 2)
    
    # Other risk factors
    data = {
        "Temperature Increase": round(np.random.uniform(0.5, 4.5) * scenario_multiplier[scenario] * time_multiplier[timeframe], 1),
        "Precipitation Change": round(np.random.uniform(-30, 20) * scenario_multiplier[scenario] * time_multiplier[timeframe], 1),
        "Sea Level Rise": round(np.random.uniform(10, 80) * scenario_multiplier[scenario] * time_multiplier[timeframe], 1),
        "Drought Risk": round(np.random.uniform(0, 100) * scenario_multiplier[scenario] * time_multiplier[timeframe] / 100, 2),
        "Flood Risk": round(np.random.uniform(0, 100) * scenario_multiplier[scenario] * time_multiplier[timeframe] / 100, 2),
        "Agricultural Yield Change": round(np.random.uniform(-30, 5) * scenario_multiplier[scenario] * time_multiplier[timeframe], 1),
        "Economic Loss (billion USD)": max(0.1, economic_loss),
        "Health Impact Index": round(np.random.uniform(0, 100) * scenario_multiplier[scenario] * time_multiplier[timeframe] / 100, 2),
        "Water Stress Index": round(np.random.uniform(0, 100) * scenario_multiplier[scenario] * time_multiplier[timeframe] / 100, 2)
    }
    
    return data

# Generate time series data for visualization
@st.cache_data
def generate_time_series(country, region, scenario, start_year=2020, end_year=2100):
    years = list(range(start_year, end_year + 1, 10))
    
    # Create a seed based on inputs for consistent random data
    seed = hash(f"{country}_{region}_{scenario}_timeseries") % 10000
    np.random.seed(seed)
    
    # Base values that change with scenario
    scenario_multiplier = {
        "RCP 2.6 (Low emissions)": 0.7,
        "RCP 4.5 (Intermediate)": 1.0,
        "RCP 8.5 (High emissions)": 1.5
    }
    
    # Calculate metrics for each year
    data = {
        "Year": years,
        "Temperature Change": [round(np.random.uniform(0.1, 0.5) * scenario_multiplier[scenario] * (year - 2020) / 10, 2) for year in years],
        "Economic Impact": [round(len(region) * 0.2 * scenario_multiplier[scenario] * (year - 2020) / 10 * (1 + np.random.normal(0, 0.1)), 2) for year in years],
        "Agricultural Yield": [round(100 - np.random.uniform(0.5, 2.0) * scenario_multiplier[scenario] * (year - 2020) / 10, 1) for year in years],
        "Water Availability": [round(100 - np.random.uniform(0.2, 1.5) * scenario_multiplier[scenario] * (year - 2020) / 10, 1) for year in years]
    }
    
    return pd.DataFrame(data)

# Generate risk data for current selection
risk_data = generate_climate_data(selected_country, selected_region, selected_timeframe, selected_scenario)
time_series_data = generate_time_series(selected_country, selected_region, selected_scenario)

# Tab 1: Risk Assessment
with tab1:
    st.markdown("<h2 class='sub-header'>📊 Climate Risk Impact Assessment</h2>", unsafe_allow_html=True)
    
    # Create a multi-column layout
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("<div class='info-box'>", unsafe_allow_html=True)
        st.markdown(f"### {selected_region}, {selected_country}")
        st.markdown(f"**Scenario:** {selected_scenario}")
        st.markdown(f"**Timeframe:** {selected_timeframe}")
        st.markdown("</div>", unsafe_allow_html=True)
        
        # Economic impact
        st.markdown("### Economic Impact")
        st.metric(
            "Estimated Annual Economic Loss", 
            f"${risk_data['Economic Loss (billion USD)']:.2f} billion",
            f"{risk_data['Economic Loss (billion USD)'] - 5:.1f}% vs baseline"
        )
        
        # Physical risk metrics
        st.markdown("### Physical Risk Metrics")
        physical_metrics = {
            "Temperature Increase": [f"{risk_data['Temperature Increase']}°C", f"+{risk_data['Temperature Increase']/2:.1f}°C"],
            "Precipitation Change": [f"{risk_data['Precipitation Change']}%", None],
            "Sea Level Rise": [f"{risk_data['Sea Level Rise']} cm", None]
        }
        
        for metric, (value, delta) in physical_metrics.items():
            st.metric(metric, value, delta)
    
    with col2:
        st.markdown("### Risk Indices")
        
        # Create gauge charts for various indices
        risk_metrics = {
            "Drought Risk": risk_data["Drought Risk"],
            "Flood Risk": risk_data["Flood Risk"],
            "Health Impact Index": risk_data["Health Impact Index"],
            "Water Stress Index": risk_data["Water Stress Index"]
        }
        
        for metric, value in risk_metrics.items():
            fig = go.Figure(go.Indicator(
                mode="gauge+number",
                value=value,
                domain={'x': [0, 1], 'y': [0, 1]},
                title={'text': metric},
                gauge={
                    'axis': {'range': [0, 1]},
                    'bar': {'color': "darkblue"},
                    'steps': [
                        {'range': [0, 0.33], 'color': "lightgreen"},
                        {'range': [0.33, 0.66], 'color': "gold"},
                        {'range': [0.66, 1], 'color': "salmon"}
                    ],
                    'threshold': {
                        'line': {'color': "red", 'width': 4},
                        'thickness': 0.75,
                        'value': 0.8
                    }
                }
            ))
            fig.update_layout(height=200, margin=dict(l=20, r=20, t=50, b=20))
            st.plotly_chart(fig, use_container_width=True)
    
    # Additional Impact Section
    st.markdown("### Sectoral Impacts")
    
    impact_cols = st.columns(3)
    
    with impact_cols[0]:
        st.markdown("#### Agricultural Impact")
        st.metric(
            "Crop Yield Change", 
            f"{risk_data['Agricultural Yield Change']}%", 
            risk_data['Agricultural Yield Change']
        )
        
        if risk_data['Agricultural Yield Change'] < -15:
            st.markdown("<div class='warning-box'>Significant food security risk.</div>", unsafe_allow_html=True)
    
    with impact_cols[1]:
        st.markdown("#### Water Resources")
        st.metric(
            "Water Stress", 
            f"{risk_data['Water Stress Index']*100:.1f}%", 
            f"{risk_data['Water Stress Index']*100 - 30:.1f}%"
        )
        
        if risk_data['Water Stress Index'] > 0.6:
            st.markdown("<div class='warning-box'>Critical water scarcity expected.</div>", unsafe_allow_html=True)
    
    with impact_cols[2]:
        st.markdown("#### Health Impacts")
        st.metric(
            "Health Risk Index", 
            f"{risk_data['Health Impact Index']*100:.1f}", 
            f"{risk_data['Health Impact Index']*100 - 25:.1f}%"
        )
        
        if risk_data['Health Impact Index'] > 0.5:
            st.markdown("<div class='warning-box'>Increased disease burden expected.</div>", unsafe_allow_html=True)

# Tab 2: Data Visualization
with tab2:
    st.markdown("<h2 class='sub-header'>📈 Climate Risk Visualizations</h2>", unsafe_allow_html=True)
    
    # Time series data visualization
    st.markdown("### Projected Changes Over Time")
    
    # Select visualization metric
    viz_metric = st.selectbox(
        "Select metric to visualize",
        ["Temperature Change", "Economic Impact", "Agricultural Yield", "Water Availability"]
    )
    
    # Create visualization
    fig = px.line(
        time_series_data, 
        x="Year", 
        y=viz_metric,
        markers=True,
        title=f"{viz_metric} Projection for {selected_region}, {selected_country} ({selected_scenario})"
    )
    
    fig.update_layout(
        xaxis_title="Year",
        yaxis_title=viz_metric,
        legend_title="Scenario",
        height=500
    )
    
    st.plotly_chart(fig, use_container_width=True)
    
    # Comparison visualization
    st.markdown("### Cross-Region Comparison")
    
    # Generate comparison data across regions
    comparison_regions = region_list[selected_country][:4] if selected_country in region_list else ["Region A", "Region B", "Region C"]
    
    comparison_data = {
        "Region": comparison_regions,
        "Economic Loss": [
            generate_climate_data(selected_country, region, selected_timeframe, selected_scenario)["Economic Loss (billion USD)"]
            for region in comparison_regions
        ],
        "Temperature Increase": [
            generate_climate_data(selected_country, region, selected_timeframe, selected_scenario)["Temperature Increase"]
            for region in comparison_regions
        ],
        "Agricultural Impact": [
            generate_climate_data(selected_country, region, selected_timeframe, selected_scenario)["Agricultural Yield Change"]
            for region in comparison_regions
        ]
    }
    
    comparison_df = pd.DataFrame(comparison_data)
    
    comp_metric = st.radio(
        "Select comparison metric", 
        ["Economic Loss", "Temperature Increase", "Agricultural Impact"],
        horizontal=True
    )
    
    # Create bar chart
    fig = px.bar(
        comparison_df,
        x="Region",
        y=comp_metric,
        title=f"{comp_metric} Comparison Across Regions ({selected_timeframe}, {selected_scenario})",
        color=comp_metric,
        color_continuous_scale="Blues" if comp_metric == "Economic Loss" else "RdYlGn_r"
    )
    
    st.plotly_chart(fig, use_container_width=True)

# Tab 3: Report Generation
with tab3:
    st.markdown("<h2 class='sub-header'>📄 Export Reports & Data</h2>", unsafe_allow_html=True)
    
    st.markdown("""
    Generate comprehensive reports of your climate risk analysis in various formats.
    These reports include all the metrics, projections, and visualizations from your selected parameters.
    """)
    
    # Report customization
    report_options = st.multiselect(
        "Select report sections to include",
        ["Executive Summary", "Economic Impacts", "Physical Risk Metrics", "Sectoral Analysis", "Time Series Projections", "Regional Comparisons", "Adaptation Recommendations"],
        default=["Executive Summary", "Economic Impacts", "Physical Risk Metrics"]
    )
    
    include_charts = st.checkbox("Include charts and visualizations", value=True)
    
    col1, col2 = st.columns(2)
    
    with col1:
        # PDF Report Generation
        st.markdown("### PDF Report")
        if st.button("Generate PDF Report"):
            try:
                # Create a temporary file
                with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
                    temp_filename = temp_file.name
                
                # Generate PDF
                c = canvas.Canvas(temp_filename, pagesize=letter)
                c.setFont("Helvetica-Bold", 16)
                c.drawString(100, 750, f"Climate Risk Assessment Report")
                c.setFont("Helvetica", 12)
                c.drawString(100, 730, f"Region: {selected_region}, {selected_country}")
                c.drawString(100, 710, f"Scenario: {selected_scenario}, Timeframe: {selected_timeframe}")
                
                # Economic impact section
                c.setFont("Helvetica-Bold", 14)
                c.drawString(100, 670, "Economic Impact Assessment")
                c.setFont("Helvetica", 12)
                c.drawString(100, 650, f"Estimated Annual Economic Loss: ${risk_data['Economic Loss (billion USD)']:.2f} billion")
                
                # Physical risks section
                c.setFont("Helvetica-Bold", 14)
                c.drawString(100, 610, "Physical Risk Metrics")
                c.setFont("Helvetica", 12)
                c.drawString(100, 590, f"Temperature Increase: {risk_data['Temperature Increase']}°C")
                c.drawString(100, 570, f"Precipitation Change: {risk_data['Precipitation Change']}%")
                c.drawString(100, 550, f"Sea Level Rise: {risk_data['Sea Level Rise']} cm")
                
                # Save the PDF
                c.save()
                
                st.markdown("<p class='success-text'>✅ PDF Report Generated Successfully!</p>", unsafe_allow_html=True)
                
                # Provide download link
                with open(temp_filename, "rb") as pdf_file:
                    st.download_button(
                        label="Download PDF Report",
                        data=pdf_file,
                        file_name=f"CRISAP_{selected_country}_{selected_region}_Report.pdf",
                        mime="application/pdf"
                    )
                
                # Clean up temporary file
                os.unlink(temp_filename)
                
            except Exception as e:
                st.error(f"Error generating PDF: {str(e)}")
    
    with col2:
        # DOCX Report Generation
        st.markdown("### DOCX Report")
        if st.button("Generate DOCX Report"):
            try:
                # Create a temporary file
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                    temp_filename = temp_file.name
                
                # Create docx
                doc = Document()
                doc.add_heading(f"Climate Risk Assessment Report", level=0)
                doc.add_heading(f"{selected_region}, {selected_country}", level=1)
                doc.add_paragraph(f"Scenario: {selected_scenario}")
                doc.add_paragraph(f"Timeframe: {selected_timeframe}")
                
                # Executive Summary
                if "Executive Summary" in report_options:
                    doc.add_heading("Executive Summary", level=1)
                    doc.add_paragraph(f"""
                    This report presents a comprehensive climate risk assessment for {selected_region}, {selected_country} 
                    under the {selected_scenario} scenario with projections to {selected_timeframe}. The analysis indicates 
                    significant economic, agricultural, and environmental impacts if current trends continue.
                    """)
                
                # Economic Impacts
                if "Economic Impacts" in report_options:
                    doc.add_heading("Economic Impact Assessment", level=1)
                    doc.add_paragraph(f"Estimated Annual Economic Loss: ${risk_data['Economic Loss (billion USD)']:.2f} billion")
                    doc.add_paragraph(f"This represents approximately {risk_data['Economic Loss (billion USD)']/5:.1f}% of the regional GDP.")
                
                # Physical Risk Metrics
                if "Physical Risk Metrics" in report_options:
                    doc.add_heading("Physical Risk Metrics", level=1)
                    doc.add_paragraph(f"Temperature Increase: {risk_data['Temperature Increase']}°C")
                    doc.add_paragraph(f"Precipitation Change: {risk_data['Precipitation Change']}%")
                    doc.add_paragraph(f"Sea Level Rise: {risk_data['Sea Level Rise']} cm")
                
                # Sectoral Analysis
                if "Sectoral Analysis" in report_options:
                    doc.add_heading("Sectoral Impact Analysis", level=1)
                    doc.add_paragraph(f"Agricultural Yield Change: {risk_data['Agricultural Yield Change']}%")
                    doc.add_paragraph(f"Water Stress Index: {risk_data['Water Stress Index']*100:.1f}%")
                    doc.add_paragraph(f"Health Impact Index: {risk_data['Health Impact Index']*100:.1f}")
                
                # Adaptation Recommendations
                if "Adaptation Recommendations" in report_options:
                    doc.add_heading("Adaptation Recommendations", level=1)
                    doc.add_paragraph("Based on the risk assessment, the following adaptation measures are recommended:")
                    recommendations = doc.add_paragraph()
                    recommendations.add_run("1. Develop drought-resistant crop varieties\n").bold = True
                    recommendations.add_run("2. Implement improved water management systems\n").bold = True
                    recommendations.add_run("3. Strengthen coastal protection infrastructure\n").bold = True
                    recommendations.add_run("4. Enhance early warning systems for extreme weather events").bold = True
                
                # Save the document
                doc.save(temp_filename)
                
                st.markdown("<p class='success-text'>✅ DOCX Report Generated Successfully!</p>", unsafe_allow_html=True)
                
                # Provide download link
                with open(temp_filename, "rb") as docx_file:
                    st.download_button(
                        label="Download DOCX Report",
                        data=docx_file,
                        file_name=f"CRISAP_{selected_country}_{selected_region}_Report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                
                # Clean up temporary file
                os.unlink(temp_filename)
                
            except Exception as e:
                st.error(f"Error generating DOCX: {str(e)}")
    
    # Data export section
    st.markdown("### Export Raw Data")
    export_format = st.radio("Select export format", ["CSV", "Excel", "JSON"], horizontal=True)
    
    # Prepare data for export
    export_data = pd.DataFrame({
        "Metric": list(risk_data.keys()),
        "Value": list(risk_data.values()),
        "Region": selected_region,
        "Country": selected_country,
        "Scenario": selected_scenario,
        "Timeframe": selected_timeframe
    })
    
    # Show data preview
    st.dataframe(export_data, use_container_width=True)
    
    # Export buttons based on format
    if export_format == "CSV":
        csv = export_data.to_csv(index=False)
        st.download_button(
            label="Download CSV",
            data=csv,
            file_name=f"CRISAP_{selected_country}_{selected_region}_Data.csv",
            mime="text/csv"
        )
    elif export_format == "Excel":
        # For Excel, we'd need additional libraries like openpyxl
        # For demonstration, we'll use CSV as a fallback
        csv = export_data.to_csv(index=False)
        st.download_button(
            label="Download as CSV (Excel format not available in this version)",
            data=csv,
            file_name=f"CRISAP_{selected_country}_{selected_region}_Data.csv",
            mime="text/csv"
        )
    elif export_format == "JSON":
        json_data = export_data.to_json(orient="records")
        st.download_button(
            label="Download JSON",
            data=json_data,
            file_name=f"CRISAP_{selected_country}_{selected_region}_Data.json",
            mime="application/json"
        )

# Footer
st.markdown("---")
st.markdown(
    """
    <div style="text-align: center; color: gray; font-size: 0.8rem;">
        CRISAP - Climate Risk and Impact Scenarios Analysis Platform | v1.0 | © 2025
    </div>
    """, 
    unsafe_allow_html=True
)
